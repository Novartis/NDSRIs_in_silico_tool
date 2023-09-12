Copyright 2021 Novartis Institutes for BioMedical Research Inc.
 
Licensed under the MIT License (the "License");
you may not use this file except in compliance with the License.
You may obtain a copy of the License at
 
https://www.mit.edu/~amini/LICENSE.md
 
Unless required by applicable law or agreed to in writing, software
distributed under the License is distributed on an "AS IS" BASIS,
WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
See the License for the specific language governing permissions and
limitations under the License.

import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import pandas as pd


#Create Download Report file from the Calcualtion
def create_word_file(image_path_1,image_path_2,message,score):
    df = pd.DataFrame(message, columns=['Column Name'])
    df_1 = pd.read_csv('rule_table.csv')
    df_2 =pd.read_csv('docx_head_table.csv',header=None)
    document = Document()
    paragraph0 = document.add_paragraph()
    paragraph0.add_run("XXX").bold=True
    # head table
    t = document.add_table(rows=(len(df_2)), cols=2)
    t.style = 'Table Grid'
    t.autofit = False
    for i in range(0,len(df_2)):
        row_cells = df_2.iloc[i]
        for j in range(len(row_cells)):
            t.cell(i, j).text = row_cells[j]

    # introduction
    document.add_paragraph('')
    para_intro = document.add_paragraph()
    run1 = para_intro.add_run("Introduction")
    run1.bold=True
    para_intro.add_run('\n')
    para_intro.add_run('(Example, please edit) A1 are impuritiy contain a N-nitroso-moiety(Figure 1)')
    para_intro.add_run('\n')
    document.add_picture(image_path_1,width=Inches(3.5),height =Inches(3.5))
    document.add_paragraph("Figure 1: Structures information")

    # safety assessment
    document.add_paragraph('')
    para_safety = document.add_paragraph()
    run2 = para_safety.add_run("Safety assessment")
    run2.bold=True
    para_safety.add_run('\n')
    para_safety.add_run("(Example,please edit) A1 was not mutagenic in the Ames test (study xxxxxxx). However, the Ames test is not fully in compliance of the enhanced Ames test (EAT) according to EMA (2023). EMA (2023) reports that a limit of 1500 ng/day can be applied for NDSRIs that are negative in the EAT.")
    para_safety.add_run('\n')
    para_safety.add_run('EMA (2023) suggests that if N-nitrosamines are identified without sufficient substance specific data to derive a substance specific limit for lifetime exposure as recommended in ICH M7(R1) guideline, the “Carcinogenic Potency Categorization Approach” (CPCA) for N-nitrosamines (Annex 2) should be used to establish the AI (Acceptable Intake) unless other robust data are available that would override this AI.')
    para_safety.add_run('\n')
    para_safety.add_run('The EMA flowchart to predict the carcinogenicity potency is provided in Figure 2. Its application to A1 is presented in Table 1.')

    # flowchart picture
    document.add_picture(image_path_2,width=Inches(3),height =Inches(4.5))
    document.add_paragraph("Figure 2: Flowchart to Predict the Potency Category of an N-nitrosamine")

    # table 1
    document.add_paragraph("Table 1 Calculating Potency Category")
    tt = document.add_table(rows=len(df_1)+1, cols=2)
    tt.style = 'Table Grid'
    tt.autofit = False
    hdr_cells = tt.rows[0].cells
    hdr_cells[0].text = 'Flowchart question'
    hdr_cells[1].text = 'Reply / Comment'
    tt.cell(0,0).paragraphs[0].runs[0].bold=True
    tt.cell(0,1).paragraphs[0].runs[0].bold=True
    for i in range(1,len(df_1)+1):
        row_cells = df_1.iloc[i-1]
        for j in range(len(row_cells)):
            tt.cell(i, j).text = row_cells[j]
            if i == 9:
                tt.cell(i, j).paragraphs[0].runs[0].bold=True

    if score == None or score == 100:
        if 'No alpha hydrogen on its alpha carbon' in message:
            # category = 'Potency Category 5 : AI = 1500 ng/day'
            AI = '1500 ng/day'
            tt.cell(9, 1).text = '1500 ng/day'
            tt.cell(1, 1).text = 'Yes'
        if 'Do not have more than one alpha hydrogen on its alpha carbons' in message:
            # category = 'Potency Category 5 : AI = 1500 ng/day'
            AI = '1500 ng/day'
            tt.cell(9, 1).text = '1500 ng/day'
            tt.cell(2, 1).text = 'Yes'
        if 'Have a tertiary alpha carbon' in message:
            # category = 'Potency Category 5 : AI = 1500 ng/day'
            AI = '1500 ng/day'
            tt.cell(9, 1).text = '1500 ng/day'
            tt.cell(3, 1).text = 'Yes'
    elif score >= 4:
        # category = 'Potency Category 4 : AI = 1500 ng/day'
        AI = '1500 ng/day'
        tt.cell(9, 1).text = '1500 ng/day'
        tt.cell(5, 1).text = 'Yes'
    elif score == 3:
        # category = 'Potency Category 3 : AI = 400 ng/day'
        AI = '400 ng/day'
        tt.cell(9, 1).text = '400 ng/day'
        tt.cell(6, 1).text = 'Yes'
    elif score == 2:
        # category = 'Potency Category 2 : AI = 100 ng/day'
        AI = '100 ng/day'
        tt.cell(9, 1).text = '100 ng/day'
        tt.cell(7, 1).text = 'Yes'
    elif score <= 1:
        # category = 'Potency Category 1 : AI = 18 ng/day'
        AI = '18 ng/day'
        tt.cell(9, 1).text = '18 ng/day'
        tt.cell(8, 1).text = 'Yes'

    document.add_paragraph('* Potency Score = α-Hydrogen Score + Deactivating Feature Score (sum all scores for features present in the N-nitrosamine) + Activating Feature Score (sum all scores for features present in the N-nitrosamine')
    document.add_paragraph("")

    # table 2
    document.add_paragraph('Table 2 Explicit Rules being Applied')
    t2 = document.add_table(rows=(len(df)+1), cols=1)
    t2.style = 'Table Grid'
    t2.autofit = False
    heading_cells = t2.rows[0].cells
    heading_cells[0].text = 'Applicable Criteria As Given below'
    heading_cells[0].paragraphs[0].runs[0].bold = True
    for i in range(0,len(df)):
        row_cells = df.iloc[i]
        t2.cell(i+1, 0).text = row_cells[0]

    # conclusion
    document.add_paragraph('')
    para_con = document.add_paragraph()
    para_con_run = para_con.add_run("Conclusion")
    para_con_run.bold=True
    para_con.add_run("\n")
    para_con.add_run(f'(Example, please edit)The AI for A1, is {AI}. N-Nitrosamines present below 10% of their respective AI do not need to be factored into the calculation of limits for individual or total N-nitrosamine(s). Otherwise, the respective N-nitrosamine should be included in the calculation of a total N-nitrosamines limit based on the most potent one (EMA, 2023).')

    # reference
    para_ref = document.add_paragraph()
    para_ref_run = para_ref.add_run("Reference")
    para_ref_run.bold = True
    para_ref.add_run("\n")
    para_ref.add_run("EMA (European Medicines Agency) (2023), 28 July 2023, EMA/409815/2020 Rev.17 Questions and answers for marketing authorisation holders/applicants on the CHMP Opinion for the Article 5(3) of Regulation (EC) No 726/2004 referral on nitrosamine impurities in human medicinal products.")

    # save doc
    bio = io.BytesIO()
    document.save(bio)
    bio.seek(0)
    st.download_button(label='Download Word File', data=bio.getvalue(), file_name='output.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
